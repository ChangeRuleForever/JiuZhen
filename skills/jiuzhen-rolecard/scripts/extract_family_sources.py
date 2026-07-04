from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


REPO_ROOT = Path(__file__).resolve().parents[3]
SOURCE_ROOT = REPO_ROOT / "source"
FAMILY_SOURCE_ROOT = SOURCE_ROOT / "families"
TABLE_SOURCE_ROOT = SOURCE_ROOT / "tables"

DOCS = [
    "【曹家】商法体系.docx",
    "乐正家.docx",
    "墨家_新.docx",
    "亓家.docx",
    "林家.docx",
    "张家.docx",
    "岳家.docx",
    "易家.docx",
    "巫家.docx",
    "药家.docx",
]


def clean(text: str) -> str:
    return " ".join((text or "").replace("\u3000", " ").split())


def read_doc(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return {
        "file": path.name,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def read_trait_xlsx(path: Path) -> dict:
    wb = load_workbook(path, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = []
    for row in ws.iter_rows(values_only=True):
        values = [clean("" if v is None else str(v)) for v in row]
        if any(values):
            rows.append(values)
    return {
        "file": path.name,
        "sheet": ws.title,
        "rows": rows,
    }


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    result = {
        "trait_xlsx": read_trait_xlsx(TABLE_SOURCE_ROOT / "通常特性体系.xlsx"),
        "docs": [read_doc(FAMILY_SOURCE_ROOT / name) for name in DOCS],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
